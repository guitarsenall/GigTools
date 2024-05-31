
'''
scratch.py
'''

import gigtools as gt
import datetime as dt


## get the venue and date from title
#title   = 'Germantown Spring Fling 5/4/24'
#tok     = title.rpartition(' ')
#venue   = tok[0]
#datestr = tok[2]
#m,d,y   = datestr.split('/')
#date    = dt.date( int(y)+2000, int(m), int(d) )
#diff    = dt.date.today() - date


# read a DOCX file containing multiple gigs using docx
import docx
import os
import unicodedata
ScoreThreshold  = 70.0
repertwaar      = gt.read_repertwaar()
GigFolder       = 'S:\\will\\documents\\OneDrive\\2024\\gigtools\\gigfiles\\'
GigFile         = 'multi.docx'
gig_files       = [ 'gigs_january_2024.docx'    ,
                    'gigs_february_2024.docx'   ,
                    'gigs_march_2024.docx'      ,
                    'gigs_april_2024.docx'      ,
                    'gigs_may_2024.docx'        ]
gigs    = []
for GigFile in gig_files:
    doc = docx.Document(GigFolder + GigFile)
    allText = []
    for docpara in doc.paragraphs:
        allText.append(unicodedata.normalize('NFKD',docpara.text))
    # find the title lines
    title_idx   = []
    for i, line in enumerate(allText):
        if 'Performance #' in line:
            #print(f'Found one at {i}: {line}')
            b   = 0
            for b in range(10):
                if i-b < 0:
                    print(f'Missing title for line {i}: {line}')
                    break
                if '/24' in allText[i-b]:
                    title_idx.append(i-b)
                    #print(f'\tTitle: {allText[i-b]}')
                    break
    print(f'Found {len(title_idx)} gig titles in {GigFile}:')
    for i in title_idx:
        print('\t' + allText[i])
    title_idx.append(len(allText)-1)
    # loop over the gig sections
    for i in range(len(title_idx)-1):
        gig_lines   = allText[ title_idx[i] : title_idx[i+1]-1 ]
        title       = allText[title_idx[i]]
        print(f'Parsing Gig title: {title}')
        # get the venue and date
        tok         = title.rpartition(' ')
        venue       = tok[0]
        datestr     = tok[2]
        m,d,y       = datestr.split('/')
        gigdate     = dt.date( int(y)+2000, int(m), int(d) )
        # get the song list
        blocks      = []
        song_list   = []
        this_block  = []
        for line in gig_lines:
            if line == ' ':
                blocks.append(this_block[:])    # copy the list
                this_block  = []
            else:
                this_block.append(line)
        for block in blocks:
            if len(block) >= 9:
                matches = 0
                for line in block:
                    score   = gt.score_line(repertwaar, line)
                    if score >= ScoreThreshold:
                        # increment the match count
                        matches += 1
                if matches / len(block) >= 0.50:
                    # song list found
                    song_list   = block
                    break
        gig_songs   = gt.match_gig_songs(repertwaar, song_list)
        print(f'\t{len(gig_songs)} songs found' )
        print(f'\t{len(gig_lines)} text lines' )
        for song in gig_songs:
            song['playcount']  += 1
            song['playdates'].append(gigdate)
        gig = { 'Title' : allText[title_idx[i]] ,
                'Venue' : venue                 ,
                'Date'  : gigdate               ,
                'Songs' : gig_songs             ,
                'Lines' : gig_lines             }
        gigs.append(gig)
played_songs    = []
for song in repertwaar:
    if song['playcount'] > 0:
        played_songs.append(song)
repertwaar.sort( reverse=True, key = lambda s: s['playcount'] )
playcount   = 0
print('play counts by song:')
for song in repertwaar:
    if song['playcount'] > 0:
        # form the play-interval string
        gig_dates   = sorted( song['playdates'], reverse=True )
        CurDate     = dt.date.today()
        delta       = CurDate - gig_dates[0]
        DeltaString = f'{delta.days}'
        CurDate     = gig_dates[0]
        for gigdate in gig_dates[1:]:
            delta       = CurDate - gigdate
            DeltaString += f',{delta.days}'
            CurDate     = gigdate
    else:
        DeltaString = ''
    print( '\t{0:<40}: {1:d} plays : {2}'.format(
                song['title'].removesuffix('*'),
                song['playcount'], DeltaString ) )



## read a DOCX gig file using docx
#import docx
#import os
#import unicodedata
#ScoreThreshold  = 70.0
#repertwaar      = gt.read_repertwaar()
#GigFolder       = 'S:\\will\\documents\\OneDrive\\2024\\gigtools\\gigfiles\\'
#GigFile         = 'Tres Rojas 91023.docx'
#doc = docx.Document(GigFolder + GigFile)
#allText = []
#for docpara in doc.paragraphs:
#    allText.append(unicodedata.normalize('NFKD',docpara.text))
#blocks  = []
#this_block  = []
#for line in allText:
#    if line == ' ':
#        blocks.append(this_block[:])    # copy the list
#        this_block  = []
#    else:
#        this_block.append(line)
#for block in blocks:
#    if len(block) >= 5:
#        matches = 0
#        for line in block:
#            score   = gt.score_line(repertwaar, line)
#            if score >= ScoreThreshold:
#                # increment the match count
#                matches += 1
#        if matches / len(block) >= 0.50:
#            # song list found
#            song_list   = block
#            break
#gig_songs   = gt.match_gig_songs(repertwaar, song_list)
#for song in gig_songs:
#    song['playcount']  += 1
#for song in repertwaar:
#    print(song['title'].removesuffix('*'), ': ',
#            str(song['playcount']), ' plays' )


## read a PDF gig file using pypdf
#from pypdf import PdfReader
#import os
#os.chdir('S:\\will\\documents\\OneDrive\\2024\\gigtools')
#reader = PdfReader('gigfiles/Tres Rojas 91023.pdf')
#print(len(reader.pages))
#page = reader.pages[0]
#text = page.extract_text()
#print(text)


## read, reverse, and print a list of songs in a file
#file = open('songlist.txt')
#lines = []
#for line in file.readlines():
#    line = line.replace('\n', '')
#    lines.append(line)
#file.close()
#songs = lines
#songs.reverse()
#print(*songs, sep='\n')


## print a simple list of songs
#songs   = [
#            "The Mission"                    ,
#            "Under The Sun"                  ,
#            "Emmaus"                         ,
#            "As Warm As Tears"               ,
#            "Those Who Wait"                 ,
#            "Del's Bells"                    ,
#            "Nothing But The Call"           ,
#            "Time In A Bottle"               ,
#            "Still The One"                  ,
#            "The Wind & The Wheat"           ,
#            "You've Got A Friend"            ,
#            "Solitary Man"                   ,
#            "Misunderstanding"               ,
#            "If You Could Read My Mind"      ,
#            "Here, There and Everywhere"     ,
#            "Crazy Little Thing Called Love" ,
#            "Even The Losers"                ,
#            "Shilo"                          ,
#            "Kashmir"                        ,
#            "Metamorphosis"                  ,
#            "Classical Gas"                  ,
#            "Doctor My Eyes"                 ,
#            "Behind Blue Eyes"               ,
#            "Addison's Walk"                 ,
#            "Beginnings"                     ,
#            "Sailing"                        ,
#            "Blackbird"                      ,
#            "Everything I Own"               ,
#            "Infinite Horizon / County Down" ,
#            "The Reunion"                    ,
#            "Wild World"                     ,
#            "Here Comes The Sun"             ,
#            "Time After Time"                ,
#            "Something"                      ,
#            "I've Got A Name"                ]
#songs.reverse()
#print(*songs, sep='\n')

